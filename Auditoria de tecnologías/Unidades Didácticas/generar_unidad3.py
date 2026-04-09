#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generador de la Unidad III - Auditoria de Tecnologias de Informacion
Archivo: AT Unidad III.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# CONFIGURACION DEL DOCUMENTO
# ============================================================
import pathlib
BASE_DIR = pathlib.Path("C:/Users/ingga/OneDrive/Documentos/Nueva carpeta/Auditoria de tecnologías/Unidades Didácticas")
RUTA = str(BASE_DIR / "AT_Unidad_III.docx")

doc = Document()

# Configuracion de pagina
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.space_before = Pt(0)
pf.line_spacing = 1.15

# ============================================================
# FUNCIONES AUXILIARES
# ============================================================

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h

def add_para(text, bold=False, italic=False, alignment=None, space_after=None, space_before=None, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_mixed_para(parts, alignment=None, space_after=None, space_before=None):
    """parts = lista de tuplas (texto, bold, italic, size, color)"""
    p = doc.add_paragraph()
    for texto, bold, italic, sz, clr in parts:
        run = p.add_run(texto)
        run.bold = bold
        run.italic = italic
        if sz:
            run.font.size = Pt(sz)
        if clr:
            run.font.color.rgb = clr
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

def add_numbered(text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

def create_table(headers, rows, col_widths=None):
    """Crea tabla con encabezado estilizado."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Encabezado
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1B3A5C")

    # Filas de datos
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EBF0F5")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1B3A5C"/></w:pBdr>')
    pPr.append(pBdr)

def add_intro_box(text):
    """Caja de introduccion con fondo azul claro."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF0F5"/>')
    pPr.append(shading)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def add_autoevaluacion(preguntas, respuestas):
    """Agrega seccion de autoevaluacion."""
    add_heading_styled("Autoevaluacion", level=2)
    add_para("Responde las siguientes preguntas para verificar tu comprension sobre los temas de esta unidad. Revisa tus apuntes y el material antes de comenzar.", italic=True, space_after=12)

    for i, (pregunta, opciones) in enumerate(preguntas, 1):
        add_para(f"Pregunta {i}:", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x5C))
        add_para(pregunta, space_after=4)
        for opt in opciones:
            add_bullet(opt)
        add_para("", space_after=6)

    add_para("Respuestas:", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x5C), space_before=12)
    for r in respuestas:
        add_para(r, space_after=2)


# ============================================================
# PORTADA
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AUDITORIA EN TECNOLOGIAS DE INFORMACION")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Unidad III")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Ejecucion de la Auditoria de TI")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

for _ in range(4):
    doc.add_paragraph()

add_para("Planeacion de la auditoria: definicion de alcance y objetivos", alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, color=RGBColor(0x55, 0x55, 0x55))
add_para("Ejecucion de la Auditoria y recoleccion de evidencias", alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, color=RGBColor(0x55, 0x55, 0x55))
add_para("Informe de Auditoria", alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, color=RGBColor(0x55, 0x55, 0x55))
add_para("Seguimiento del informe", alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, color=RGBColor(0x55, 0x55, 0x55))

doc.add_page_break()


# ============================================================
# INDICE DE CONTENIDO
# ============================================================
add_heading_styled("Indice de Contenido", level=1)
add_separator()

indice_items = [
    ("Introduccion", "3"),
    ("Desarrollo de Contenidos", "4"),
    ("  3.1. Planeacion de la auditoria: definicion de alcance y objetivos", "4"),
    ("  3.2. Ejecucion de la Auditoria y recoleccion de evidencias", "15"),
    ("  3.3. Informe de Auditoria", "26"),
    ("  3.4. Seguimiento del informe", "36"),
    ("Evaluacion Integral de la Unidad III", "44"),
    ("Bibliografia y Webgrafia", "52"),
    ("Glosario", "52"),
]

for item, page in indice_items:
    p = doc.add_paragraph()
    if item.startswith("  "):
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(item.strip())
        run.font.size = Pt(11)
    else:
        run = p.add_run(item)
        run.bold = True
        run.font.size = Pt(11)
    tab = p.add_run("\t" + page)
    tab.font.size = Pt(11)

doc.add_page_break()


# ============================================================
# INTRODUCCION
# ============================================================
add_heading_styled("Introduccion", level=1)
add_separator()

add_intro_box(
    "En las unidades anteriores estudiamos los fundamentos conceptuales de la auditoria de TI "
    "y los marcos metodologicos que guian su practica. Ahora damos el paso a la accion concreta: "
    "como se planifica, ejecuta, comunica y da seguimiento a una auditoria de TI en el mundo real."
)

intro_parrafos = [
    "La auditoria en Tecnologias de Informacion no es un ejercicio teorico; es un proceso practico "
    "y estructurado que requiere disciplina, metodo y rigor profesional. De nada sirve conocer los "
    "mejores marcos de referencia (COBIT, ITIL, ISO 27001) si no sabemos como aplicarlos paso a paso "
    "en una auditoria real.",

    "Esta Unidad III, titulada \"Ejecucion de la Auditoria de TI\", se centra en el ciclo de vida "
    "completo del proceso de auditoria. Abordaremos las cuatro fases operativas fundamentales: la "
    "planeacion, la ejecucion, la comunicacion de resultados a traves del informe y el seguimiento "
    "de las recomendaciones emitidas.",

    "A lo largo de cuatro temas, desarrollaremos cada una de estas fases con detalle, proporcionando "
    "herramientas practicas, ejemplos contextualizados en el entorno nicaraguense y ejercicios de "
    "autoevaluacion que te permitiran verificar tu comprension."
]

for par in intro_parrafos:
    add_para(par)

add_para("En el Tema 3.1", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x5C))
add_para("profundizaremos en la planeacion de la auditoria, la fase mas critica del proceso. Aprenderas "
    "a definir el alcance y los objetivos de manera precisa, a realizar una evaluacion preliminar de "
    "riesgos, a seleccionar los criterios de auditoria apropiados y a elaborar un programa de trabajo "
    "detallado. Veremos que una buena planificacion es la diferencia entre una auditoria exitosa y una "
    "que desperdicia recursos sin generar valor.")

add_para("En el Tema 3.2", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x5C))
add_para("entraremos en el trabajo de campo: la ejecucion de la auditoria y la recoleccion de evidencias. "
    "Exploraremos las tecnicas de auditoria (entrevistas, revision documental, observacion, cuestionarios), "
    "los tipos de pruebas (cumplimiento y sustantivas), las herramientas de auditoria asistida por "
    "computadora (CAATTs) y la importancia de documentar todo en los papeles de trabajo.")

add_para("En el Tema 3.3", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x5C))
add_para("aprenderas a comunicar los resultados de la auditoria a traves del informe. Descubriras como "
    "estructurar un hallazgo de auditoria con sus cinco atributos (condicion, criterio, causa, efecto y "
    "recomendacion), como clasificar los hallazgos por severidad, como redactar conclusiones y como "
    "elaborar un informe profesional que sea util para la toma de decisiones.")

add_para("En el Tema 3.4", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x5C))
add_para("cerraremos el ciclo con el seguimiento del informe. Comprenderas por que la auditoria no termina "
    "con la entrega del informe, como obtener y monitorear el plan de accion de la gerencia, como verificar "
    "que las recomendaciones se implementaron y como el seguimiento genera valor real para la organizacion.")

add_intro_box(
    "Al finalizar esta unidad, estaras en capacidad de planificar, ejecutar, informar y dar seguimiento "
    "a una auditoria de TI de manera profesional, estructurada y alineada con los estandares internacionales."
)

doc.add_page_break()


# ============================================================
# DESARROLLO DE CONTENIDOS
# ============================================================
add_heading_styled("Desarrollo de Contenidos", level=1)
add_separator()


# ============================================================
# TEMA 3.1: PLANEACION DE LA AUDITORIA
# ============================================================
add_heading_styled("3.1. Planeacion de la auditoria: definicion de alcance y objetivos", level=2)
add_separator()

add_para(
    "En la Unidad II, estudiamos el proceso metodologico de la auditoria de TI y sus cinco fases. "
    "Ahora profundizaremos en cada una de ellas, comenzando por la mas importante: la planeacion. "
    "Si hay una regla de oro en auditoria, es que una mala planificacion conduce inevitablemente a "
    "una auditoria ineficaz, sin importar cuan competente sea el auditor."
)

add_para("Imagina que un medico intenta diagnosticar a un paciente sin conocer su historial, sin saber que "
    "sintomas tiene y sin un plan de examenes. Probablemente ordenaria pruebas innecesarias, pasaria por "
    "alto la causa real del problema y desperdiciaria tiempo y recursos. Lo mismo ocurre con una auditoria "
    "de TI sin planificacion adecuada.", italic=True)

add_heading_styled("La importancia de la planeacion en la auditoria de TI", level=3)

add_para(
    "La planeacion es la fase donde el auditor \"disena\" la auditoria. Es el momento de pensar antes "
    "de actuar, de definir que se va a hacer, como, cuando y con que recursos. Una planificacion "
    "rigurosa aporta beneficios tangibles que impactan toda la auditoria:"
)

tabla_beneficios = create_table(
    ["Beneficio", "Descripcion", "Consecuencia de NO planificar"],
    [
        ["Rigor y sistematicidad",
         "Asegura que la auditoria cubra todos los aspectos importantes de manera ordenada y completa.",
         "La auditoria se vuelve erratic; el auditor salta de un tema a otro sin un hilo conductor."],
        ["Eficiencia en el uso de recursos",
         "Guia al auditor, evitando que invierta tiempo en areas irrelevantes o que pase por alto areas criticas.",
         "Se desperdician horas en revisar aspectos menores mientras los riesgos graves quedan sin evaluar."],
        ["Objetividad",
         "Proporciona un marco que reduce la subjetividad y los sesgos personales del auditor.",
         "Los hallazgos pueden ser cuestionados por basarse en opiniones personales en lugar de un metodo estructurado."],
        ["Defendibilidad",
         "Los hallazgos y conclusiones, al basarse en un plan reconocido, son mas faciles de defender ante la gerencia o un comite.",
         "La gerencia puede rechazar los hallazgos alegando que la auditoria fue arbitraria o incompleta."],
        ["Calidad",
         "Facilita la supervision y revision del trabajo, asegurando el cumplimiento de los estandares de calidad.",
         "No hay forma de verificar si la auditoria se realizo con el nivel de calidad esperado."],
    ]
)

add_para("")
add_para(
    "En el contexto nicaraguense, donde muchas organizaciones tienen recursos limitados para TI y "
    "para auditoria, una buena planificacion es aun mas crucial. No se puede permitir el lujo de "
    "desperdiciar tiempo y esfuerzo en areas que no generan valor."
)

add_heading_styled("Actividades preliminares de planeacion", level=3)

add_para("Antes de definir alcance y objetivos, el auditor debe realizar una serie de actividades "
    "preliminares que le proporcionen el contexto necesario para tomar decisiones informadas.")

add_heading_styled("1. Conocimiento del negocio y del entorno", level=4)

add_para(
    "El auditor no puede auditar lo que no comprende. Antes de tocar un solo sistema, debe entender "
    "la organizacion en su conjunto. Esto implica:"
)

conocimiento_items = [
    ("Mision, vision y objetivos estrategicos:", "Que hace la organizacion, hacia donde va y "
     "como la tecnologia apoya esos objetivos. Por ejemplo, si auditas el sistema de banca en linea "
     "de BANPRO, debes entender que la mision del banco es ofrecer servicios financieros confiables "
     "y que la banca en linea es un canal critico para esa mision."),
    ("Procesos de negocio principales:", "Cuales son los procesos core del negocio y cuales sistemas "
     "de TI los soportan. En una empresa como La Colonia, el proceso de gestion de inventarios y "
     "facturacion es core, y su sistema ERP es fundamental."),
    ("Estructura organizativa del area de TI:", "Cuantas personas trabajan en TI, cual es su "
     "estructura jerarquica, quien toma las decisiones tecnologicas."),
    ("Marco legal y normativo aplicable:", "Que leyes, regulaciones y normas aplican a la "
     "organizacion. Para una entidad publica como la Alcaldia de Managua, aplican las NAGUN. "
     "Para un banco, las normas de la SIBOIF. Para quien procesa pagos con tarjeta, PCI DSS. "
     "Para todos, la Ley 1042 de Proteccion de Datos Personales."),
    ("Auditorias anteriores:", "Revisar los informes de auditorias previas (internas y externas) "
     "para identificar hallazgos pendientes y patrones recurrentes."),
]

for titulo, desc in conocimiento_items:
    add_mixed_para([
        (titulo, True, False, None, None),
        (desc, False, False, None, None),
    ], space_after=6)

add_heading_styled("2. Reunion de apertura con el auditado", level=4)

add_para(
    "Una vez que el auditor ha recopilado la informacion preliminar, debe reunirse con los "
    "responsables del area o sistema a auditar. Esta reunion tiene varios propositos:"
)

add_bullet("Presentar al equipo de auditoria y establecer una relacion de trabajo basada en la confianza.")
add_bullet("Confirmar los objetivos generales de la auditoria y asegurarse de que el auditado los entiende.")
add_bullet("Solicitar la documentacion inicial necesaria (politicas, procedimientos, organigramas, informes previos).")
add_bullet("Establecer un cronograma tentativo y acordar las fechas de las actividades.")
add_bullet("Resolver dudas iniciales y identificar a las personas clave que seran entrevistadas.")

add_para(
    "Esta reunion no es una confrontacion. El auditor debe transmitir un mensaje claro: la auditoria "
    "no viene a buscar culpables, sino a evaluar controles y proponer mejoras. Un auditado que entiende "
    "esto colaborara mucho mas que uno que se siente bajo sospecha.", italic=True
)

add_heading_styled("Evaluacion preliminar de riesgos", level=3)

add_para(
    "Con el conocimiento del negocio y del entorno, el auditor puede ahora identificar los riesgos "
    "que podrian afectar al area o sistema auditado. Esta evaluacion preliminar es fundamental porque "
    "determina donde se enfocaran los esfuerzos de la auditoria. No se puede auditar todo con el mismo "
    "nivel de detalle; hay que priorizar."
)

add_heading_styled("Identificacion de activos criticos", level=4)

add_para("El primer paso es identificar los activos de TI mas importantes para la organizacion:")

add_bullet("Activos de hardware: servidores, routers, firewalls, equipos de almacenamiento.")
add_bullet("Activos de software: sistemas ERP, core bancario, aplicaciones criticas, bases de datos.")
add_bullet("Activos de informacion: datos de clientes, informacion financiera, propiedad intelectual, datos personales.")
add_bullet("Recursos humanos: administradores de sistemas, DBA, personal clave con conocimientos especializados.")

add_heading_styled("Identificacion de amenazas y vulnerabilidades", level=4)

add_para("Para cada activo critico, el auditor identifica las amenazas potenciales y las vulnerabilidades "
    "existentes:")

add_mixed_para([
    ("Amenazas: ", True, False, None, None),
    ("Eventos o acciones que podrian causar dano. Ejemplos: ciberataques (ransomware, phishing), "
     "fallos de hardware, errores humanos, desastres naturales (en Nicaragua, terremotos y huracanes "
     "son amenazas reales), fraudes internos.", False, False, None, None),
], space_after=6)

add_mixed_para([
    ("Vulnerabilidades: ", True, False, None, None),
    ("Debilidades que podrian ser explotadas por las amenazas. Ejemplos: servidores sin parches de "
     "seguridad, contraseñas por defecto, falta de respaldo de datos, ausencia de politicas de seguridad, "
     "personal sin capacitacion.", False, False, None, None),
], space_after=6)

add_heading_styled("Matriz de riesgos preliminar", level=4)

add_para(
    "Con las amenazas y vulnerabilidades identificadas, el auditor construye una matriz de riesgos "
    "preliminar, evaluando cada riesgo en terminos de probabilidad e impacto:"
)

tabla_matriz = create_table(
    ["Riesgo identificado", "Probabilidad\n(Alta/Media/Baja)", "Impacto\n(Alto/Medio/Bajo)", "Prioridad"],
    [
        ["Acceso no autorizado al sistema de facturacion por cuentas de exempleados activas",
         "Alta", "Alto", "ALTA"],
        ["Falta de backups verificables del core bancario",
         "Media", "Alto", "ALTA"],
        ["Servidores de desarrollo sin parches de seguridad",
         "Media", "Medio", "MEDIA"],
        ["Ausencia de plan de continuidad documentado para el sistema de nomina",
         "Baja", "Alto", "MEDIA"],
        ["WiFi de visitantes en la misma red que servidores internos",
         "Alta", "Medio", "ALTA"],
    ]
)

add_para("")
add_para(
    "Esta matriz preliminar le dice al auditor donde enfocar sus recursos. Los riesgos de prioridad "
    "ALTA seran los objetivos principales de la auditoria."
)

add_heading_styled("Definicion del alcance", level=3)

add_para(
    "El alcance de la auditoria define los limites de la revision. Establece que se incluye y que "
    "se excluye, evitando malentendidos y expectativas irreales."
)

add_para("Un alcance bien definido debe especificar:", bold=True, space_after=6)

alcance_items = [
    "Sistemas y aplicaciones: Que sistemas especificos seran auditados (ej. \"modulo de facturacion del SAP\", \"sistema de nomina\").",
    "Infraestructura: Que servidores, bases de datos, dispositivos de red estan dentro del alcance.",
    "Procesos: Que procesos de TI se evaluaran (ej. \"gestion de cambios\", \"gestion de accesos\", \"gestion de backups\").",
    "Ubicaciones fisicas: Que oficinas, data centers o sucursales se visitaran (ej. \"oficina central de Managua y centro de datos principal\").",
    "Periodo de tiempo: Que periodo cubre la auditoria (ej. \"transacciones y registros del ano 2025\").",
    "Exclusiones explicitas: Que NO se incluira (ej. \"no se auditaran las sucursales departamentales\", \"no se realizaran pruebas de penetracion\").",
]

for item in alcance_items:
    add_bullet(item)

add_para("Ejemplo de definicion de alcance contextualizado:", bold=True, space_after=6)

add_intro_box(
    "\"La auditoria abarcara el modulo de cuentas por cobrar del sistema contable de la Alcaldia de Leon, "
    "los servidores de base de datos MySQL que lo soportan, las politicas de gestion de accesos y gestion "
    "de cambios aplicables a dicho sistema, y las operaciones realizadas durante el periodo enero-diciembre "
    "2025. La auditoria se limitara a las instalaciones de la oficina central de la Alcaldia. No se "
    "incluiran las terminales de las ventanillas de cobro ni las sucursales municipales.\""
)

add_heading_styled("Definicion de objetivos", level=3)

add_para(
    "Los objetivos definen que se espera lograr con la auditoria. Deben ser claros, medibles y "
    "alineados con los riesgos identificados en la evaluacion preliminar."
)

add_para("Los objetivos se formulan en dos niveles:", bold=True, space_after=6)

add_mixed_para([
    ("Objetivo general: ", True, False, None, None),
    ("La meta amplia de la auditoria, que responde a la pregunta \"que queremos lograr con esta auditoria\".",
     False, False, None, None),
], space_after=4)

add_mixed_para([
    ("Objetivos especificos: ", True, False, None, None),
    ("Las metas concretas y medibles que desglosan el objetivo general. Cada objetivo especifico "
     "debe poder evaluarse como cumplido o no cumplido al final de la auditoria.",
     False, False, None, None),
], space_after=6)

add_para("Ejemplo aplicado a una auditoria en TELCOR:", bold=True, space_after=6)

tabla_objetivos = create_table(
    ["Tipo", "Objetivo"],
    [
        ["General",
         "Evaluar la efectividad de los controles de seguridad logica y fisica sobre el sistema de gestion "
         "de nombres de dominio .ni de TELCOR."],
        ["Especifico 1",
         "Verificar si los controles de gestion de accesos garantizan que solo personal autorizado puede "
         "modificar registros de dominios."],
        ["Especifico 2",
         "Determinar si la seguridad fisica del data center protege adecuadamente los servidores del sistema."],
        ["Especifico 3",
         "Evaluar si existen y se cumplen politicas de gestion de cambios para las modificaciones al sistema de dominios."],
        ["Especifico 4",
         "Verificar si se realizan respaldos periodicamente y si se ha probado su restauracion."],
    ]
)

add_para("")

add_heading_styled("Seleccion de criterios de auditoria", level=3)

add_para(
    "Los criterios de auditoria son las normas, estandares, politicas o requisitos contra los cuales "
    "el auditor comparara la evidencia que recolecte. Sin criterios, no hay forma de determinar si algo "
    "esta \"bien\" o \"mal\"."
)

add_para("La seleccion de criterios depende del tipo de auditoria y del contexto:", bold=True, space_after=6)

tabla_criterios = create_table(
    ["Contexto de la auditoria", "Criterios aplicables"],
    [
        ["Auditoria en entidad publica (ministerio, alcaldia)",
         "NAGUN (Normas de Auditoria Gubernamental de Nicaragua), politicas internas de TI de la entidad, "
         "Ley 1042 de Proteccion de Datos Personales."],
        ["Auditoria en banco o entidad financiera",
         "Normas de la SIBOIF sobre gestion de riesgos de TI, COBIT, ISO 27001, PCI DSS (si procesa tarjetas), "
         "politicas internas del banco."],
        ["Auditoria de seguridad de la informacion",
         "ISO/IEC 27001:2022 (Anexo A), NIST Cybersecurity Framework, politicas de seguridad de la organizacion."],
        ["Auditoria de cumplimiento PCI DSS",
         "Los 12 requisitos del estandar PCI DSS v4.0."],
        ["Auditoria de gestion de servicios de TI",
         "ITIL 4 (practicas de gestion de incidentes, cambios, niveles de servicio), acuerdos de nivel de "
         "servicio (SLA) internos."],
    ]
)

add_para("")

add_heading_styled("Elaboracion del programa de trabajo", level=3)

add_para(
    "El programa de trabajo es la \"hoja de ruta\" detallada de la auditoria. Es el documento que "
    "lista, en orden, todos los procedimientos que se realizaran, con su responsable y tiempo estimado. "
    "Es la materializacion de toda la planificacion."
)

add_para("Un programa de trabajo debe incluir:", bold=True, space_after=6)

add_bullet("Numero de referencia de cada procedimiento.")
add_bullet("Descripcion clara del procedimiento a realizar.")
add_bullet("Objetivo del procedimiento (que se busca verificar).")
add_bullet("Criterio de auditoria asociado.")
add_bullet("Responsable de ejecutar el procedimiento.")
bullet_tiempo = add_bullet("Tiempo estimado de dedicacion.")
add_bullet("Fecha programada de ejecucion.")
add_bullet("Espacio para registrar los resultados y referencias a los papeles de trabajo.")

add_para("Ejemplo de programa de trabajo:", bold=True, space_after=6)

tabla_programa = create_table(
    ["No.", "Procedimiento", "Objetivo", "Responsable", "Tiempo"],
    [
        ["P-01", "Entrevista con el administrador del sistema de facturacion",
         "Comprender el flujo de procesamiento y los controles de acceso implementados.",
         "Auditor A", "2 horas"],
        ["P-02", "Revision de la politica de gestion de accesos",
         "Verificar que exista una politica documentada y actualizada.",
         "Auditor A", "1 hora"],
        ["P-03", "Prueba de cumplimiento: revision de 30 solicitudes de acceso",
         "Verificar que todas las cuentas de usuario tienen autorizacion documentada.",
         "Auditor B", "4 horas"],
        ["P-04", "Prueba sustantiva: analisis de logs de acceso fallidos del ultimo trimestre",
         "Identificar patrones de intentos de acceso no autorizado.",
         "Auditor B", "3 horas"],
        ["P-05", "Observacion directa: proceso de backup del servidor",
         "Verificar que el backup se realiza segun el procedimiento documentado.",
         "Auditor A", "2 horas"],
        ["P-06", "Prueba de recorrido (walkthrough): procesamiento de una factura desde su creacion hasta su registro contable",
         "Confirmar el entendimiento del flujo de datos y los controles aplicados.",
         "Auditor A y B", "3 horas"],
    ]
)

add_para("")
add_para(
    "El programa de trabajo es flexible. Si durante la ejecucion el auditor descubre riesgos no "
    "identificados en la planificacion, debe poder ajustar el programa para incluir procedimientos "
    "adicionales. Pero cualquier cambio debe estar justificado y documentado."
)

add_heading_styled("Carta de encargo (memorando de inicio)", level=3)

add_para(
    "La carta de encargo es la comunicacion formal que el auditor envia al auditado para informarle "
    "sobre la auditoria. Establece el tono de la relacion y asegura que todos los involucrados "
    "comprendan el proposito y el alcance."
)

add_para("Contenido tipico de una carta de encargo:", bold=True, space_after=6)

add_bullet("Destinatario (generalmente el gerente del area auditada o el titular de la entidad).")
add_bullet("Referencia al mandato que origina la auditoria (plan anual de auditoria, solicitud de la alta direccion, etc.).")
add_bullet("Objetivos de la auditoria.")
add_bullet("Alcance (sistemas, procesos, periodo, ubicaciones).")
add_bullet("Criterios de auditoria que se utilizaran.")
add_bullet("Equipo de auditoria (nombres y cargos).")
add_bullet("Cronograma tentativo (fechas de inicio, trabajo de campo, reunion de cierre).")
add_bullet("Solicitud de designar un enlace o punto focal dentro del area auditada.")
add_bullet("Declaracion de confidencialidad.")

add_para("", space_after=6)

add_heading_styled("Resumen del Tema 3.1", level=3)

add_para(
    "La planeacion es la piedra angular de toda auditoria de TI exitosa. En esta fase, el auditor "
    "conoce el negocio, evalua los riesgos preliminares, define con precision que se va a auditar "
    "(alcance) y para que (objetivos), selecciona los criterios contra los cuales comparara la "
    "realidad y elabora un programa de trabajo detallado que servira como hoja de ruta."
)

add_para(
    "Un auditor que planifica bien tiene muchas mas probabilidades de identificar los riesgos reales, "
    "recolectar evidencia relevante, emitir hallazgos utiles y generar valor para la organizacion. "
    "Un auditor que planifica mal, por muy competente que sea, terminara con una auditoria superficial, "
    "ineficiente y cuestionable."
)

add_para(
    "En el contexto nicaraguense, donde los recursos son frecuentemente limitados y las organizaciones "
    "necesitan maximizar el valor de cada auditoria, una planeacion rigurosa no es un lujo: es una "
    "necesidad.", italic=True
)

add_separator()
add_autoevaluacion(
    [
        ("Cual es la principal razon por la que la planeacion se considera la fase mas importante de la auditoria?", [
            "a) Porque es la fase que mas tiempo consume.",
            "b) Porque una mala planificacion conduce a una auditoria ineficaz, independientemente de la competencia del auditor.",
            "c) Porque en esta fase se redacta el informe final.",
            "d) Porque es la unica fase donde se recolecta evidencia."
        ]),
        ("Un auditor esta revisando un sistema de gestion tributaria de una alcaldia nicaraguense y decide que los criterios de auditoria seran las NAGUN y las politicas internas de la entidad. Esta actividad corresponde a:", [
            "a) Definicion del alcance.",
            "b) Evaluacion preliminar de riesgos.",
            "c) Seleccion de criterios de auditoria.",
            "d) Elaboracion del programa de trabajo."
        ]),
        ("Cual de los siguientes elementos NO forma parte de un alcance de auditoria bien definido?", [
            "a) Los sistemas y aplicaciones que se auditaran.",
            "b) El periodo de tiempo que cubre la auditoria.",
            "c) Las recomendaciones que se emitiran.",
            "d) Las exclusiones explicitas (lo que NO se incluira)."
        ]),
        ("En una matriz de riesgos preliminar, un riesgo con probabilidad \"Alta\" e impacto \"Alto\" debe clasificarse con prioridad:", [
            "a) Baja.",
            "b) Media.",
            "c) Alta.",
            "d) Se excluye de la auditoria."
        ]),
        ("Un programa de trabajo de auditoria debe incluir todos los siguientes elementos, EXCEPTO:", [
            "a) Descripcion de los procedimientos a realizar.",
            "b) Responsable de cada procedimiento.",
            "c) El salario del auditor asignado.",
            "d) Tiempo estimado de dedicacion."
        ]),
        ("Caso practico: Eres el auditor lider de una firma contratada por BANPRO para auditar su sistema de banca en linea. Durante la evaluacion preliminar de riesgos, identificas que el mayor riesgo es la posibilidad de accesos no autorizados desde internet. Como definiras el enfoque y los objetivos especificos de la auditoria en respuesta a este riesgo?", [
            "a) Auditar todos los sistemas de BANPRO sin excepcion.",
            "b) Enfocar la auditoria en los controles de seguridad logica del sistema de banca en linea (autenticacion, firewalls, registros de acceso) y definir objetivos especificos para verificar cada uno de estos controles.",
            "c) Cambiar el alcance para auditar solo la seguridad fisica del data center.",
            "d) No incluir el riesgo en el programa de trabajo porque es demasiado tecnico."
        ]),
    ],
    [
        "b) Porque una mala planificacion conduce a una auditoria ineficaz, independientemente de la competencia del auditor.",
        "c) Seleccion de criterios de auditoria.",
        "c) Las recomendaciones que se emitiran (las recomendaciones se formulan al final, no en el alcance).",
        "c) Alta.",
        "c) El salario del auditor asignado (no es un componente del programa de trabajo).",
        "b) Enfocar la auditoria en los controles de seguridad logica del sistema de banca en linea y definir objetivos especificos para verificar cada uno de estos controles.",
    ]
)

doc.add_page_break()


print("Tema 3.1 completado. Guardando progreso...")
doc.save(RUTA)
print(f"Progreso guardado en: {RUTA}")
